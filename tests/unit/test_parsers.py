from pathlib import Path

from docx import Document as DocxDocument
from reportlab.pdfgen import canvas

from local_agentic_rag.parsers import parse_document


def test_docx_parser_extracts_heading_sections(tmp_path: Path) -> None:
    docx_path = tmp_path / "playbook.docx"
    document = DocxDocument()
    document.add_heading("Playbook", level=1)
    document.add_heading("Checklist", level=2)
    document.add_paragraph("First response within 4 business hours.")
    document.save(docx_path)

    parsed = parse_document(docx_path)
    assert parsed.detected_title == "Playbook"
    assert parsed.sections[0].section_path.endswith("Checklist")


def test_pdf_parser_extracts_page_sections(tmp_path: Path) -> None:
    pdf_path = tmp_path / "notes.pdf"
    pdf = canvas.Canvas(str(pdf_path))
    pdf.drawString(100, 750, "Pricing Notes")
    pdf.drawString(100, 735, "Annual plan includes a 12 percent discount.")
    pdf.save()

    parsed = parse_document(pdf_path)
    assert parsed.detected_title == "Pricing Notes"
    assert parsed.sections[0].page_number == 1
