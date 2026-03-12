from __future__ import annotations

import re
from pathlib import Path

from .models import ParsedDocument, ParsedSection


CONTENT_TYPES = {
    ".pdf": "application/pdf",
    ".md": "text/markdown",
    ".markdown": "text/markdown",
    ".txt": "text/plain",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def parse_document(path: Path) -> ParsedDocument:
    suffix = path.suffix.lower()
    if suffix in {".md", ".markdown"}:
        return _parse_markdown(path)
    if suffix == ".txt":
        return _parse_text(path)
    if suffix == ".docx":
        return _parse_docx(path)
    if suffix == ".pdf":
        return _parse_pdf(path)
    raise ValueError(f"Unsupported file type: {path.suffix}")


def _parse_markdown(path: Path) -> ParsedDocument:
    raw = path.read_text(encoding="utf-8")
    lines = raw.splitlines()
    heading_stack: list[str] = []
    buffer: list[str] = []
    sections: list[ParsedSection] = []
    section_start = 1
    detected_title = path.stem.replace("_", " ").title()

    def flush(current_line: int) -> None:
        if not buffer:
            return
        text = "\n".join(buffer).strip()
        buffer.clear()
        if not text:
            return
        section_name = " > ".join(heading_stack) if heading_stack else detected_title
        sections.append(
            ParsedSection(
                text=text,
                section_path=section_name,
                line_start=section_start,
                line_end=max(section_start, current_line - 1),
            )
        )

    for index, line in enumerate(lines, start=1):
        match = re.match(r"^(#{1,6})\s+(.+?)\s*$", line)
        if match:
            flush(index)
            heading_level = len(match.group(1))
            heading_text = match.group(2).strip()
            if heading_level == 1 and detected_title == path.stem.replace("_", " ").title():
                detected_title = heading_text
            heading_stack[:] = heading_stack[: heading_level - 1] + [heading_text]
            section_start = index + 1
            continue
        buffer.append(line)
    flush(len(lines) + 1)

    if not sections and raw.strip():
        sections.append(ParsedSection(text=raw.strip(), section_path=detected_title, line_start=1, line_end=len(lines)))
    return ParsedDocument(
        source_path=str(path),
        content_type=CONTENT_TYPES[path.suffix.lower()],
        detected_title=detected_title,
        sections=sections,
    )


def _parse_text(path: Path) -> ParsedDocument:
    raw = path.read_text(encoding="utf-8")
    lines = raw.splitlines()
    non_empty = [line.strip() for line in lines if line.strip()]
    detected_title = non_empty[0] if non_empty and len(non_empty[0]) < 90 else path.stem.replace("_", " ").title()
    paragraphs = re.split(r"\n\s*\n", raw)
    sections: list[ParsedSection] = []
    cursor = 1
    for paragraph in paragraphs:
        text = paragraph.strip()
        line_count = max(1, paragraph.count("\n") + 1)
        if text:
            sections.append(
                ParsedSection(
                    text=text,
                    section_path=detected_title,
                    line_start=cursor,
                    line_end=cursor + line_count - 1,
                )
            )
        cursor += line_count + 1
    return ParsedDocument(
        source_path=str(path),
        content_type=CONTENT_TYPES[path.suffix.lower()],
        detected_title=detected_title,
        sections=sections,
    )


def _parse_docx(path: Path) -> ParsedDocument:
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover - handled in runtime docs
        raise RuntimeError("python-docx is required for DOCX support.") from exc

    document = DocxDocument(str(path))
    heading_stack: list[str] = []
    sections: list[ParsedSection] = []
    buffer: list[str] = []
    section_start = 1
    detected_title = path.stem.replace("_", " ").title()

    def flush(current_index: int) -> None:
        if not buffer:
            return
        text = "\n".join(buffer).strip()
        buffer.clear()
        if not text:
            return
        section_name = " > ".join(heading_stack) if heading_stack else detected_title
        sections.append(
            ParsedSection(
                text=text,
                section_path=section_name,
                line_start=section_start,
                line_end=max(section_start, current_index - 1),
            )
        )

    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if not text:
            continue
        style_name = getattr(paragraph.style, "name", "") or ""
        if style_name.lower().startswith("heading"):
            flush(index)
            heading_level = 1
            parts = style_name.split()
            if len(parts) > 1 and parts[-1].isdigit():
                heading_level = int(parts[-1])
            heading_stack[:] = heading_stack[: heading_level - 1] + [text]
            if heading_level == 1 and detected_title == path.stem.replace("_", " ").title():
                detected_title = text
            section_start = index + 1
            continue
        buffer.append(text)
    flush(len(document.paragraphs) + 1)

    return ParsedDocument(
        source_path=str(path),
        content_type=CONTENT_TYPES[path.suffix.lower()],
        detected_title=detected_title,
        sections=sections,
    )


def _parse_pdf(path: Path) -> ParsedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - handled in runtime docs
        raise RuntimeError("pypdf is required for PDF support.") from exc

    reader = PdfReader(str(path))
    sections: list[ParsedSection] = []
    detected_title = path.stem.replace("_", " ").title()

    for page_index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if page_index == 1 and lines and len(lines[0]) < 100:
            detected_title = lines[0]
        paragraphs = [paragraph.strip() for paragraph in re.split(r"\n\s*\n", text) if paragraph.strip()]
        line_cursor = 1
        for paragraph in paragraphs:
            line_count = max(1, paragraph.count("\n") + 1)
            sections.append(
                ParsedSection(
                    text=paragraph,
                    section_path=f"Page {page_index}",
                    page_number=page_index,
                    line_start=line_cursor,
                    line_end=line_cursor + line_count - 1,
                )
            )
            line_cursor += line_count + 1
    return ParsedDocument(
        source_path=str(path),
        content_type=CONTENT_TYPES[path.suffix.lower()],
        detected_title=detected_title,
        sections=sections,
    )
